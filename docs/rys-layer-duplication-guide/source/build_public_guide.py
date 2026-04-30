#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/home/benbi/qwen36_rys_blog")
SOURCE = ROOT / "qwen36_rys_public_guide_source.txt"
DOCX_OUT = ROOT / "rys_layer_duplication_guide.docx"
MD_OUT = ROOT / "post_qwen36_rys_no_looping.md"


VISUALS = {
    "VISUAL_SLOT_INTRO_HERO": ("assets/intro_loaded_not_correct_generated.png", "The cold-open problem: a model can load cleanly while the copied layer story is still wrong."),
    "VISUAL_SLOT_TWO_LEDGERS_STORY": ("assets/two_ledgers_same_map_generated.png", "The tensor ledger and config ledger both need to follow the same map."),
    "VISUAL_SLOT_WRONG_TRACK": ("assets/wrong_track_mismatch_generated.png", "The bug is not that one attention type is bad. The bug is sending copied weights down the wrong execution path."),
    "VISUAL_SLOT_QUANT_CALLBACK": ("assets/quantization_callback_generated.png", "BF16 is the first audition. Quantization is the callback."),
    "VISUAL_SLOT_LAYER_PACKAGE": ("assets/layer_assumptions_travel_together.png", "A copied layer has weights and execution assumptions. Move both."),
    "VISUAL_SLOT_BUILD_ROUTE": ("assets/rys_build_route_map.png", "The build route: map, tensors, config, validation, testing, quantization."),
    "VISUAL_SLOT_BENCH_OBJECTS": ("assets/rys_objects_on_the_bench.png", "Before editing a checkpoint, keep the five objects separate: weights, names, config, per-layer metadata, and the converted artifact."),
    "VISUAL_SLOT_RUNTIME_VS_PROVENANCE": ("assets/runtime_order_vs_source_provenance.png", "Takeaway: the model still runs forward; the map only records where each output layer came from."),
    "VISUAL_SLOT_HALF_OPEN_RANGE": ("assets/half_open_range_number_line.png", "Half-open ranges: blocks:15,20 copies 15..19, not 20."),
    "VISUAL_SLOT_TOY_MAPPING": ("assets/toy_mapping_blocks_2_4.png", "Toy mapping for blocks:2,4."),
    "VISUAL_SLOT_INSERT_MAPPING": ("assets/rys_insert_mapping_blocks_15_20.png", "Real mapping for blocks:15,20: 64 layers become 69."),
    "VISUAL_SLOT_OFFSET_MATH": ("assets/offset_shift_after_insert.png", "The offset is just the insert length: d = j - i."),
    "VISUAL_SLOT_TENSOR_RENAME": ("assets/tensor_rename_conveyor.png", "Tensor values can be copied, but output layer keys must be renamed."),
    "VISUAL_SLOT_CONFIG_MIRROR": ("assets/config_mirror_table.png", "Per-layer config lists must follow the same output-to-source map."),
    "VISUAL_SLOT_LAYER_TYPES": ("assets/layer_types_is_execution_plan.png", "In Qwen3.6 hybrid models, layer_types is an execution plan."),
    "VISUAL_SLOT_LOAD_PATHS": ("assets/loader_paths_safetensors_vs_gguf.png", "HF-folder loading and GGUF loading validate different artifacts."),
    "VISUAL_SLOT_VALIDATION_CARD": ("assets/validation_checklist_card.png", "Mechanical checks before quality judgment."),
    "VISUAL_SLOT_TESTING_LADDER": ("assets/testing_ladder.png", "Test like a ladder: structure first, quantization last."),
    "VISUAL_SLOT_WINDOW_FUNNEL": ("assets/window_selection_funnel.png", "Window selection belongs to the exact source checkpoint."),
    "VISUAL_SLOT_SCOREBOARD_COMPARISON": ("assets/same_suite_scoreboard.png", "Compare candidates on the same suite. Different probes are useful notes, not one shared leaderboard."),
    "VISUAL_SLOT_QUANT_CARDS": ("assets/quantization_survival_cards.png", "Quantization is a second trial, not a formality."),
    "VISUAL_SLOT_RESULTS_CHART": ("assets/strict_probe_results_chart.png", "Strict probe results for the Qwen3.6 AEON source line."),
    "VISUAL_SLOT_MASTER_MAP": ("assets/master_map_summary.png", "One output-to-source map feeds the whole build: tensors, config metadata, validation, testing, and quantization."),
}


HEADINGS = {
    "The model loaded. That was the trap.",
    "What RYS is actually changing",
    "Before cutting anything, map the machine",
    "The map is the story",
    "Where the offset actually comes from",
    "One map, two ledgers",
    "The Qwen3.6 loop story",
    "Loaded where?",
    "The checks that saved time",
    "Test like a ladder",
    "A practical validation path",
    "Choosing the layer window",
    "Quantization is the callback",
    "The Qwen3.6 case file",
    "The finished mental model",
}


CALLOUT_STARTS = (
    "Every output layer must",
    "loaded is not the same",
    "The useful definition is simple",
    "The five objects on the bench",
    "The map is ancestry, not runtime control",
    "HF safetensors loading checks",
    "The practical rule is simple",
    "What success looks like",
    "At this point,",
    "BF16 is the audition",
    "Do not compare headline means",
)


def source_lines() -> list[str]:
    return SOURCE.read_text(encoding="utf-8").splitlines()


def is_codeish(line: str) -> bool:
    stripped = line.strip()
    return (
        "->" in stripped
        or stripped.startswith("mapping[")
        or stripped.startswith("if out")
        or stripped.startswith("i =")
        or stripped.startswith("j =")
        or stripped.startswith("duplicate_length =")
        or stripped.startswith("new_list[")
        or stripped.startswith("weights_out[")
        or stripped.startswith("layer_types_out[")
        or stripped.startswith("text_config.")
        or stripped.startswith("for ")
        or stripped.startswith("src =")
        or stripped.startswith("new_name =")
        or stripped.startswith("output_tensors[")
        or stripped.startswith("assert ")
        or stripped.startswith("combined =")
        or bool(re.match(r"^(source|output):", stripped))
    )


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.6)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(20)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(28)
    styles["Title"].font.bold = True


def add_callout(document: Document, text: str) -> None:
    para = document.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.18)
    para.paragraph_format.right_indent = Inches(0.18)
    para.paragraph_format.space_before = Pt(5)
    para.paragraph_format.space_after = Pt(7)
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)


def build_docx(lines: list[str]) -> None:
    document = Document()
    style_document(document)

    title = lines[0].strip()
    document.add_paragraph(title, style="Title")
    subtitle = document.add_paragraph("A field-note style guide to RYS layer duplication, hybrid metadata, and why a model can load while still being wrong.")
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.color.rgb = RGBColor(90, 90, 90)

    for raw in lines[1:]:
        line = raw.rstrip()
        if not line:
            continue
        if line in VISUALS:
            rel_path, caption = VISUALS[line]
            image_path = ROOT / rel_path
            if image_path.exists():
                para = document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(str(image_path), width=Inches(6.55))
                cap = document.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.color.rgb = RGBColor(90, 90, 90)
            continue
        if line in HEADINGS:
            document.add_paragraph(line, style="Heading 1")
            continue
        if line.startswith(CALLOUT_STARTS):
            add_callout(document, line)
            continue
        if is_codeish(line):
            para = document.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            run = para.add_run(line)
            run.font.name = "Aptos Mono"
            run.font.size = Pt(9.5)
            continue
        document.add_paragraph(line)

    document.save(DOCX_OUT)
    (ROOT / "qwen36_rys_guide.docx").write_bytes(DOCX_OUT.read_bytes())


def build_markdown(lines: list[str]) -> None:
    out: list[str] = [
        "---",
        'title: "The model loaded. That was the trap: RYS, Qwen3.6, and layer_types"',
        "date: 2026-04-29",
        "tags: [llm, qwen, rys, transformers, quantization]",
        "draft: true",
        "---",
        "",
    ]
    title = lines[0].strip()
    out.append(f"# {title}")
    out.append("")

    in_code = False
    for raw in lines[1:]:
        line = raw.rstrip()
        if not line:
            if in_code:
                out.append("```")
                in_code = False
            out.append("")
            continue
        if line in VISUALS:
            if in_code:
                out.append("```")
                in_code = False
            rel_path, caption = VISUALS[line]
            out.append(f"![{caption}]({rel_path})")
            out.append("")
            continue
        if line in HEADINGS:
            if in_code:
                out.append("```")
                in_code = False
            out.append(f"## {line}")
            out.append("")
            continue
        if is_codeish(line):
            if not in_code:
                out.append("```text")
                in_code = True
            out.append(line)
            continue
        if in_code:
            out.append("```")
            in_code = False
            out.append("")
        if line.startswith(CALLOUT_STARTS):
            out.append(f"> **{line}**")
        else:
            out.append(line)
        out.append("")
    if in_code:
        out.append("```")
    MD_OUT.write_text("\n".join(out).rstrip() + "\n", encoding="utf-8")


def main() -> None:
    lines = source_lines()
    build_docx(lines)
    build_markdown(lines)
    print(f"wrote {DOCX_OUT}")
    print(f"wrote {MD_OUT}")


if __name__ == "__main__":
    main()
